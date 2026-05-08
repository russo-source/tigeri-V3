"""Extract plain text from uploaded PDF / Word documents.

Used by the chat upload endpoint so the orchestrator can read attachments
inline as user message context.

Hardening notes (audit-driven, 2026-04-28):
  - The bytes are passed in-memory; no filesystem path is ever touched, so
    classical path-traversal LFI doesn't apply. We still sanitize the
    user-supplied filename through ``os.path.basename`` + a printable-char
    filter so adversarial filenames can't poison logs or error responses.
  - DOCX is a ZIP archive parsed by python-docx (lxml under the hood). We
    cap PDF page count and DOCX paragraph count so a malicious archive
    can't exhaust memory by claiming a million sections.
  - lxml from 4.x onward disables DTD loading and external entity resolution
    by default — XXE is mitigated as long as we don't manually re-enable
    them, which we don't.
  - MAX_BYTES caps the upload at 10 MB; MAX_TEXT_CHARS caps decoded output
    to 30k chars. Both bound the work the parser will do.
"""

from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass


SUPPORTED_MEDIA_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "application/msword",  # .doc (treated as best-effort)
    "text/plain",
    "text/markdown",
}

MAX_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_TEXT_CHARS = 30_000  # Truncate so we don't blow the LLM context
MAX_PDF_PAGES = 200  # Bounds memory amplification on adversarial PDFs
MAX_DOCX_PARAGRAPHS = 5_000


# Filenames may only contain printable ASCII + a small set of safe characters.
# Anything outside this is replaced with ``_``. This matches the OWASP
# guidance for unsafe filename handling.
_SAFE_FILENAME_RE = re.compile(r"[^A-Za-z0-9._\- ]")


def _sanitize_filename(raw: str) -> str:
    """Strip path components and unsafe characters from a user-supplied name."""
    if not raw:
        return "upload"
    # Discard any directory components the client tried to embed.
    base = os.path.basename(raw.replace("\\", "/"))
    # Reject null bytes outright — they break log parsing.
    base = base.replace("\x00", "")
    # Replace anything outside the allow-list.
    safe = _SAFE_FILENAME_RE.sub("_", base).strip(". ")
    return safe[:255] or "upload"


class UnsupportedFileTypeError(Exception):
    pass


class FileTooLargeError(Exception):
    pass


@dataclass
class ExtractedFile:
    filename: str
    media_type: str
    text: str
    truncated: bool


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    pages = reader.pages
    if len(pages) > MAX_PDF_PAGES:
        # Bound the work; admins still get the first N pages of a long doc.
        pages = pages[:MAX_PDF_PAGES]
    for i, page in enumerate(pages):
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            parts.append(f"[page {i + 1}: extraction failed]")
    return "\n\n".join(p.strip() for p in parts if p.strip())


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    paragraphs_seen = 0
    for p in doc.paragraphs:
        if paragraphs_seen >= MAX_DOCX_PARAGRAPHS:
            parts.append("[truncated: document exceeded paragraph cap]")
            break
        if p.text and p.text.strip():
            parts.append(p.text.strip())
        paragraphs_seen += 1
    for table in doc.tables:
        if paragraphs_seen >= MAX_DOCX_PARAGRAPHS:
            break
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text.strip():
                parts.append(row_text)
                paragraphs_seen += 1
                if paragraphs_seen >= MAX_DOCX_PARAGRAPHS:
                    break
    return "\n".join(parts)


def extract(filename: str, media_type: str, data: bytes) -> ExtractedFile:
    if len(data) > MAX_BYTES:
        raise FileTooLargeError(
            f"file is {len(data)} bytes; max allowed is {MAX_BYTES}"
        )

    safe_name = _sanitize_filename(filename)

    media = (media_type or "").lower().strip() or "application/octet-stream"
    name_lower = safe_name.lower()

    # Some browsers send odd MIME types; sniff by extension as fallback.
    if media not in SUPPORTED_MEDIA_TYPES:
        if name_lower.endswith(".pdf"):
            media = "application/pdf"
        elif name_lower.endswith(".docx"):
            media = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        elif name_lower.endswith(".doc"):
            media = "application/msword"
        elif name_lower.endswith((".txt", ".md")):
            media = "text/plain"
        else:
            raise UnsupportedFileTypeError(
                f"unsupported media_type {media_type!r} for {safe_name!r}"
            )

    if media == "application/pdf":
        text = _extract_pdf(data)
    elif media in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }:
        text = _extract_docx(data)
    else:
        text = data.decode("utf-8", errors="replace")

    truncated = len(text) > MAX_TEXT_CHARS
    if truncated:
        text = text[:MAX_TEXT_CHARS] + "\n\n[…truncated]"

    return ExtractedFile(
        filename=safe_name, media_type=media, text=text, truncated=truncated
    )
